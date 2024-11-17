from flask import Flask, request, send_file, jsonify, make_response
from flask_cors import CORS
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Twips
from config import Config
import json
import io
import os
from werkzeug.utils import secure_filename

app = Flask(__name__)
app.config.from_object(Config)
CORS(app)  # Enable CORS for cross-origin requests

# Create uploads folder if it doesn't exist
os.makedirs(app.config['UPLOAD_FOLDER'], exist_ok=True)

def allowed_file(filename):
    return '.' in filename and \
           filename.rsplit('.', 1)[1].lower() in app.config['ALLOWED_EXTENSIONS']

def transform_questions(data):
    """Transform the input JSON structure to the required format"""
    transformed_questions = []
    
    # If data is already a list of questions
    questions = data if isinstance(data, list) else data.get('questions', [])
    
    if not questions:
        raise ValueError("No questions found in input data")
        
    for question in questions:
        transformed_question = {
            'text': question['text'],
            'options': [option['text'] for option in question['options']]
        }
        transformed_questions.append(transformed_question)
        
    return transformed_questions

def add_questions_to_template(template_doc, questions):
    doc = Document(template_doc)
    
    # Add a section break before questions
    doc.add_paragraph().add_run().add_break()
    
    # Add Questions section header
    heading = doc.add_paragraph("Questions")
    heading.style = 'Heading 1'
    
    # Add table for each question
    for idx, question in enumerate(questions, 1):
        # Create a table with 2 columns (question and marks)
        table = doc.add_table(rows=1, cols=2)
        table.allow_autofit = True
        table.autofit = True
        
        # Set table width to page width
        table.columns[0].width = Inches(5.5)  # Adjust this value based on your needs
        table.columns[1].width = Inches(1)    # Width for marks column
        
        # Get cells
        question_cell = table.cell(0, 0)
        marks_cell = table.cell(0, 1)
        
        # Add question number and text
        question_para = question_cell.paragraphs[0]
        question_para.add_run(f"{idx}. ").bold = True
        question_para.add_run(question['text'])
        
        # Add marks (hardcoded to 5 marks, adjust as needed)
        marks_para = marks_cell.paragraphs[0]
        marks_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        marks_para.add_run("[5]").bold = True
        
        # Add options if present
        if 'options' in question:
            for opt_idx, option in enumerate(question['options'], 1):
                option_para = question_cell.add_paragraph()
                option_para.paragraph_format.left_indent = Inches(0.5)
                option_para.add_run(f"{opt_idx}. {option}")
        
        # Add space after table
        doc.add_paragraph()
    
    return doc

@app.route('/health', methods=['GET'])
def health_check():
    return jsonify({'status': 'healthy'}), 200

@app.route('/generate-document', methods=['POST'])
def generate_document():
    try:
        # Debug: Print request files and form data
        print("Files in request:", request.files)
        print("Form data in request:", request.form)
        
        # Check if the post request has the file part
        if 'template' not in request.files:
            print("No template file in request.files")
            return jsonify({'error': 'No template file provided'}), 400
        
        template_file = request.files['template']
        print("Template filename:", template_file.filename)
        
        # Check if a file was actually selected
        if template_file.filename == '':
            print("No selected filename")
            return jsonify({'error': 'No selected file'}), 400
        
        # Validate file type
        if not allowed_file(template_file.filename):
            print(f"Invalid file type: {template_file.filename}")
            return jsonify({'error': 'Invalid file type. Only .docx files are allowed'}), 400
        
        # Check for questions data
        if not request.form.get('questions'):
            print("No questions data provided")
            return jsonify({'error': 'No questions provided'}), 400
        
        try:
            # Parse and validate questions
            input_data = json.loads(request.form.get('questions'))
            print("Questions data:", input_data)
            transformed_questions = transform_questions(input_data)
        except json.JSONDecodeError as e:
            print(f"JSON decode error: {str(e)}")
            return jsonify({'error': 'Invalid JSON format for questions'}), 400
        except ValueError as e:
            print(f"Value error: {str(e)}")
            return jsonify({'error': str(e)}), 400
        
        # Generate document
        doc = add_questions_to_template(template_file, transformed_questions)
        
        # Save to buffer
        doc_buffer = io.BytesIO()
        doc.save(doc_buffer)
        doc_buffer.seek(0)
        
        # Create response
        response = make_response(doc_buffer.getvalue())
        response.headers['Content-Type'] = 'application/vnd.openxmlformats-officedocument.wordprocessingml.document'
        response.headers['Content-Disposition'] = f'attachment; filename={secure_filename("question_paper.docx")}'
        response.headers['Access-Control-Expose-Headers'] = 'Content-Disposition'
        
        return response
        
    except Exception as e:
        print(f"Unexpected error: {str(e)}")
        return jsonify({'error': str(e)}), 500

if __name__ == '__main__':
    app.run(debug=True, host='0.0.0.0', port=5000)
# from flask import Flask, request, send_file, jsonify
# from docx import Document
# from docx.shared import Inches
# from config import Config
# import json
# import io
# import os
# from werkzeug.utils import secure_filename

# app = Flask(__name__)
# app.config.from_object(Config)

# # Create uploads folder if it doesn't exist
# os.makedirs(app.config['UPLOAD_FOLDER'], exist_ok=True)

# def allowed_file(filename):
#     return '.' in filename and \
#            filename.rsplit('.', 1)[1].lower() in app.config['ALLOWED_EXTENSIONS']

# def transform_questions(data):
#     """Transform the input JSON structure to the required format"""
#     transformed_questions = []
    
#     # If data is already a list of questions
#     questions = data if isinstance(data, list) else data.get('questions', [])
    
#     if not questions:
#         raise ValueError("No questions found in input data")
        
#     for question in questions:
#         transformed_question = {
#             'text': question['text'],
#             'options': [option['text'] for option in question['options']]
#         }
#         transformed_questions.append(transformed_question)
        
#     return transformed_questions

# def add_questions_to_template(template_doc, questions):
#     doc = Document(template_doc)
    
#     # Add a section break before questions
#     doc.add_paragraph().add_run().add_break()
    
#     # Add Questions section header
#     heading = doc.add_paragraph("Questions")
#     heading.style = 'Heading 1'
    
#     # Add each question
#     for idx, question in enumerate(questions, 1):
#         # Add question text
#         question_para = doc.add_paragraph()
#         question_para.add_run(f"{idx}. ").bold = True
#         question_para.add_run(question['text'])
        
#         # Add options
#         if 'options' in question:
#             for option in question['options']:
#                 option_para = doc.add_paragraph()
#                 option_para.paragraph_format.left_indent = Inches(0.5)
#                 option_para.add_run(f"□ {option}")  # Square checkbox
        
#         # Add space between questions
#         doc.add_paragraph()
    
#     return doc

# @app.route('/health', methods=['GET'])
# def health_check():
#     return jsonify({'status': 'healthy'}), 200

# @app.route('/generate-document', methods=['POST'])
# def generate_document():
#     try:
#         # Check if template file is provided
#         if 'template' not in request.files:
#             return jsonify({'error': 'No template file provided'}), 400
        
#         template_file = request.files['template']
        
#         # Check if template file is valid
#         if template_file.filename == '':
#             return jsonify({'error': 'No selected file'}), 400
        
#         if not allowed_file(template_file.filename):
#             return jsonify({'error': 'Invalid file type. Only .docx files are allowed'}), 400
        
#         # Check if questions are provided in request body
#         if not request.form.get('questions'):
#             return jsonify({'error': 'No questions provided'}), 400
        
#         try:
#             # Parse questions from JSON
#             input_data = json.loads(request.form.get('questions'))
#             # Transform questions to required format
#             transformed_questions = transform_questions(input_data)
#         except json.JSONDecodeError:
#             return jsonify({'error': 'Invalid JSON format for questions'}), 400
#         except ValueError as e:
#             return jsonify({'error': str(e)}), 400
        
#         # Generate document with questions
#         doc = add_questions_to_template(template_file, transformed_questions)
        
#         # Save to memory buffer
#         doc_buffer = io.BytesIO()
#         doc.save(doc_buffer)
#         doc_buffer.seek(0)
        
#         # Send file back to client
#         return send_file(
#             doc_buffer,
#             mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document',
#             as_attachment=True,
#             download_name='generated_document.docx'
#         )
        
#     except Exception as e:
#         return jsonify({'error': str(e)}), 500

# if __name__ == '__main__':
#     app.run(debug=True, host='0.0.0.0', port=5000)
